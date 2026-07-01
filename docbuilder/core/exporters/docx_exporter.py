"""
Implementação do exportador de formato Word (DOCX).
Converte a AST serializada em JSON para um arquivo DOCX aplicando fontes, margens, cores, cabeçalhos, rodapés e numeração.
"""

from pathlib import Path
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docbuilder.core.domain.entities import TemplateStyle
from docbuilder.core.domain.document_ast import (
    HeadingBlock,
    ParagraphBlock,
    ListBlock,
    TableBlock,
    ImageBlock,
    PageBreakBlock,
    SectionBreakBlock,
)
from docbuilder.core.builders.document_builder import ASTSerializer
from docbuilder.core.domain.interfaces import IExporter


class DocxExporter(IExporter):
    """
    Exportador encarregado de criar arquivos DOCX profissionais a partir da AST em JSON.
    """

    def get_supported_format(self) -> str:
        return "docx"

    def export(
        self, source_document_path: Path, output_path: Path, template: TemplateStyle
    ) -> None:
        # 1. Carrega a AST do arquivo temporário
        with open(source_document_path, "r", encoding="utf-8") as f:
            json_ast = f.read()
        blocks = ASTSerializer.deserialize_from_json(json_ast)

        # 2. Inicializa o documento Word
        doc = DocxDocument()

        # 3. Configura Margens padrão da primeira seção (seção inicial de capa/documento)
        section = doc.sections[0]
        self._apply_section_margins(section, template)

        # Configura cor principal (convertendo string hexadecimal para RGBColor)
        primary_color_rgb = self._hex_to_rgb(template.primary_color)

        # 4. Processa os blocos AST
        for idx, block in enumerate(blocks):
            if isinstance(block, SectionBreakBlock):
                # Se não for a primeira seção, cria uma nova quebra de seção física
                if idx > 0:
                    new_section = doc.add_section()
                    self._apply_section_margins(new_section, template)
                    self._apply_header_footer(new_section, template, block.title)
                else:
                    self._apply_header_footer(section, template, block.title)

            elif isinstance(block, HeadingBlock):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.keep_with_next = True

                # Mapeamento de níveis de heading
                run = p.add_run(block.text)
                run.bold = True
                run.font.name = template.font_family

                # Tamanho e cores conforme nível do título
                if block.level == 1:
                    run.font.size = Pt(template.font_size + 8)
                    run.font.color.rgb = primary_color_rgb
                elif block.level == 2:
                    run.font.size = Pt(template.font_size + 5)
                    run.font.color.rgb = primary_color_rgb
                elif block.level == 3:
                    run.font.size = Pt(template.font_size + 3)
                    run.font.color.rgb = self._hex_to_rgb(template.secondary_color)
                else:
                    run.font.size = Pt(template.font_size + 1)

            elif isinstance(block, ParagraphBlock):
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.15

                for run_data in block.runs:
                    run = p.add_run(run_data.text)
                    run.font.name = template.font_family
                    run.font.size = Pt(template.font_size)
                    run.bold = run_data.bold
                    run.italic = run_data.italic
                    run.underline = run_data.underline
                    if run_data.link_url:
                        # Estiliza o link visualmente no Word
                        run.font.color.rgb = RGBColor(59, 130, 246)
                        run.underline = True

            elif isinstance(block, ListBlock):
                style_name = "List Number" if block.ordered else "List Bullet"
                for item in block.items:
                    p = doc.add_paragraph(style=style_name)
                    p.paragraph_format.space_after = Pt(3)

                    for run_data in item.runs:
                        run = p.add_run(run_data.text)
                        run.font.name = template.font_family
                        run.font.size = Pt(template.font_size)
                        run.bold = run_data.bold
                        run.italic = run_data.italic
                        run.underline = run_data.underline

            elif isinstance(block, TableBlock):
                # Cria a tabela no Word
                num_cols = len(table_block_headers := block.headers)
                num_rows = len(block.rows) + 1  # Headers + Rows

                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = True

                # Escreve os cabeçalhos
                hdr_cells = table.rows[0].cells
                for i, col_header in enumerate(table_block_headers):
                    hdr_cells[i].text = "".join(r.text for r in col_header.runs)
                    self._style_table_cell(
                        hdr_cells[i],
                        template.primary_color,
                        is_header=True,
                        font_family=template.font_family,
                        font_size=template.font_size,
                    )

                # Escreve as linhas de dados
                for r_idx, row_cells_data in enumerate(block.rows):
                    row_cells = table.rows[r_idx + 1].cells
                    # Alternância de cores (zebra striping) para acabamento premium
                    bg_color = "#F3F4F6" if r_idx % 2 == 1 else "#FFFFFF"
                    for c_idx, cell_data in enumerate(row_cells_data):
                        row_cells[c_idx].text = "".join(r.text for r in cell_data.runs)
                        self._style_table_cell(
                            row_cells[c_idx],
                            bg_color,
                            is_header=False,
                            font_family=template.font_family,
                            font_size=template.font_size,
                        )

                # Adiciona espaçamento após a tabela
                spacer = doc.add_paragraph()
                spacer.paragraph_format.space_before = Pt(6)

            elif isinstance(block, ImageBlock):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(12)

                try:
                    p.add_run().add_picture(block.image_path, width=Inches(5.0))
                    if block.caption:
                        caption_p = doc.add_paragraph()
                        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption_p.paragraph_format.space_after = Pt(12)
                        c_run = caption_p.add_run(f"Figura: {block.caption}")
                        c_run.italic = True
                        c_run.font.size = Pt(template.font_size - 2)
                        c_run.font.name = template.font_family
                except Exception as e:
                    # Se falhar ao carregar imagem, insere aviso textual para segurança
                    err_run = p.add_run(
                        f"[Erro ao carregar imagem: {block.image_path} - {e}]"
                    )
                    err_run.bold = True
                    err_run.font.color.rgb = RGBColor(220, 38, 38)

            elif isinstance(block, PageBreakBlock):
                doc.add_page_break()

        # Salva o arquivo gerado
        doc.save(output_path)

    def _apply_section_margins(self, section, template: TemplateStyle) -> None:
        """Aplica as margens definidas no template à seção do Word."""
        section.top_margin = Inches(template.margin_top / 2.54)
        section.bottom_margin = Inches(template.margin_bottom / 2.54)
        section.left_margin = Inches(template.margin_left / 2.54)
        section.right_margin = Inches(template.margin_right / 2.54)

    def _apply_header_footer(
        self, section, template: TemplateStyle, section_title: str
    ) -> None:
        """Aplica os cabeçalhos e rodapés configurados no template à seção."""
        # Configurar cabeçalho
        header = section.header
        header_p = header.paragraphs[0]
        header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_p.text = ""  # Limpa padrão

        # Substitui variáveis dinâmicas no cabeçalho
        text_header = template.header_text.replace("#SECTION_TITLE#", section_title)
        h_run = header_p.add_run(text_header)
        h_run.font.name = template.font_family
        h_run.font.size = Pt(9)
        h_run.font.color.rgb = RGBColor(107, 114, 128)  # Cinza claro de suporte

        # Configurar rodapé
        footer = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.text = ""  # Limpa padrão

        text_footer_parts = template.footer_text.split("#PAGE#")

        if len(text_footer_parts) > 1:
            # Se o rodapé tem a variável #PAGE#, montamos dinamicamente usando a API oxml de campos de numeração
            footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Adiciona o texto antes da página
            prefix_run = footer_p.add_run(text_footer_parts[0])
            prefix_run.font.name = template.font_family
            prefix_run.font.size = Pt(9)
            prefix_run.font.color.rgb = RGBColor(107, 114, 128)

            # Insere campo dinâmico de número de página PAGE
            page_run = footer_p.add_run()
            self._add_page_number_field(page_run)
            page_run.font.name = template.font_family
            page_run.font.size = Pt(9)
            page_run.font.color.rgb = RGBColor(107, 114, 128)

            # Adiciona o texto pós página (Ex: / Total)
            suffix_text = text_footer_parts[1].replace(
                "#TOTAL#", ""
            )  # Total dinâmico requer NUMPAGES
            suffix_run = footer_p.add_run(suffix_text)
            suffix_run.font.name = template.font_family
            suffix_run.font.size = Pt(9)
            suffix_run.font.color.rgb = RGBColor(107, 114, 128)
        else:
            footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            f_run = footer_p.add_run(template.footer_text)
            f_run.font.name = template.font_family
            f_run.font.size = Pt(9)
            f_run.font.color.rgb = RGBColor(107, 114, 128)

    def _add_page_number_field(self, run) -> None:
        """Adiciona o campo XML nativo do Word PAGE para números de página dinâmicos."""
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "separate")
        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(qn("w:fldCharType"), "end")

        r = run._r
        r.append(fldChar1)
        r.append(instrText)
        r.append(fldChar2)
        r.append(fldChar3)

    def _style_table_cell(
        self, cell, bg_hex_color: str, is_header: bool, font_family: str, font_size: int
    ) -> None:
        """Aplica estilos visuais específicos (cores e fontes) às células de tabela do Word."""
        # 1. Configura background color usando XML da célula
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{bg_hex_color.replace("#", "")}"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading)

        # 2. Configura margens internas da célula (padding) para acabamento premium
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement("w:tcMar")
        for margin in ["top", "bottom", "left", "right"]:
            node = OxmlElement(f"w:{margin}")
            node.set(
                qn("w:w"), "120" if margin in ["top", "bottom"] else "180"
            )  # tamanho do padding
            node.set(qn("w:type"), "dxa")
            tcMar.append(node)
        tcPr.append(tcMar)

        # 3. Formata os parágrafos internos
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            for run in p.runs:
                run.font.name = font_family
                run.font.size = Pt(font_size - 1)
                if is_header:
                    run.bold = True
                    run.font.color.rgb = RGBColor(
                        255, 255, 255
                    )  # Letras brancas no cabeçalho
                else:
                    run.font.color.rgb = RGBColor(31, 41, 55)

    def _hex_to_rgb(self, hex_str: str) -> RGBColor:
        hex_str = hex_str.lstrip("#")
        return RGBColor(*(int(hex_str[i : i + 2], 16) for i in (0, 2, 4)))
