"""
Implementação central do compilador de documentos (DocumentBuilder).
Lê arquivos em diversos formatos de entrada, converte em blocos AST, gera o sumário e serializa em JSON intermediário.
"""

import json
import subprocess
import tempfile
from pathlib import Path
from typing import List, Dict, Any
from bs4 import BeautifulSoup
import markdown
import docx
from docbuilder.core.domain.entities import Project, TemplateStyle
from docbuilder.core.domain.document_ast import (
    Block,
    HeadingBlock,
    ParagraphBlock,
    TextRun,
    ListBlock,
    ListItem,
    TableBlock,
    TableCell,
    ImageBlock,
    PageBreakBlock,
    SectionBreakBlock,
)
from docbuilder.core.domain.interfaces import IDocumentBuilder


class ASTSerializer:
    """Classe auxiliar para serializar e deserializar a AST de blocos em JSON."""

    @staticmethod
    def block_to_dict(block: Block) -> Dict[str, Any]:
        if isinstance(block, HeadingBlock):
            return {"type": "heading", "text": block.text, "level": block.level}
        elif isinstance(block, ParagraphBlock):
            return {
                "type": "paragraph",
                "runs": [
                    {
                        "text": r.text,
                        "bold": r.bold,
                        "italic": r.italic,
                        "underline": r.underline,
                        "link_url": r.link_url,
                    }
                    for r in block.runs
                ],
            }
        elif isinstance(block, ListBlock):
            items_data = []
            for item in block.items:
                items_data.append(
                    [
                        {
                            "text": r.text,
                            "bold": r.bold,
                            "italic": r.italic,
                            "underline": r.underline,
                            "link_url": r.link_url,
                        }
                        for r in item.runs
                    ]
                )
            return {"type": "list", "items": items_data, "ordered": block.ordered}
        elif isinstance(block, TableBlock):
            headers_data = [
                [
                    {
                        "text": r.text,
                        "bold": r.bold,
                        "italic": r.italic,
                        "underline": r.underline,
                    }
                    for r in cell.runs
                ]
                for cell in block.headers
            ]
            rows_data = []
            for row in block.rows:
                rows_data.append(
                    [
                        [
                            {
                                "text": r.text,
                                "bold": r.bold,
                                "italic": r.italic,
                                "underline": r.underline,
                            }
                            for r in cell.runs
                        ]
                        for cell in row
                    ]
                )
            return {"type": "table", "headers": headers_data, "rows": rows_data}
        elif isinstance(block, ImageBlock):
            return {
                "type": "image",
                "image_path": block.image_path,
                "caption": block.caption,
            }
        elif isinstance(block, PageBreakBlock):
            return {"type": "page_break"}
        elif isinstance(block, SectionBreakBlock):
            return {"type": "section_break", "title": block.title}
        raise ValueError(f"Tipo de bloco desconhecido para serialização: {type(block)}")

    @classmethod
    def serialize_to_json(cls, blocks: List[Block]) -> str:
        return json.dumps(
            [cls.block_to_dict(b) for b in blocks], ensure_ascii=False, indent=2
        )

    @staticmethod
    def dict_to_block(data: Dict[str, Any]) -> Block:
        b_type = data.get("type")
        if b_type == "heading":
            return HeadingBlock(text=data["text"], level=data["level"])
        elif b_type == "paragraph":
            runs = [TextRun(**r) for r in data["runs"]]
            return ParagraphBlock(runs=runs)
        elif b_type == "list":
            items = []
            for item_data in data["items"]:
                runs = [TextRun(**r) for r in item_data]
                items.append(ListItem(runs=runs))
            return ListBlock(items=items, ordered=data["ordered"])
        elif b_type == "table":
            headers = [
                TableCell(runs=[TextRun(**r) for r in cell_data])
                for cell_data in data["headers"]
            ]
            rows = []
            for row_data in data["rows"]:
                row = [
                    TableCell(runs=[TextRun(**r) for r in cell_data])
                    for cell_data in row_data
                ]
                rows.append(row)
            return TableBlock(headers=headers, rows=rows)
        elif b_type == "image":
            return ImageBlock(
                image_path=data["image_path"], caption=data.get("caption")
            )
        elif b_type == "page_break":
            return PageBreakBlock()
        elif b_type == "section_break":
            return SectionBreakBlock(title=data["title"])
        raise ValueError(f"Tipo de bloco desconhecido para deserialização: {b_type}")

    @classmethod
    def deserialize_from_json(cls, json_str: str) -> List[Block]:
        data_list = json.loads(json_str)
        return [cls.dict_to_block(d) for d in data_list]


class DocumentBuilder(IDocumentBuilder):
    """
    Builder que consome a hierarquia de documentos e gera um manifesto AST unificado em JSON.
    """

    def build(self, project: Project, base_path: Path, template: TemplateStyle) -> Path:
        all_blocks: List[Block] = []

        # 1. Capa / Título do Projeto
        all_blocks.append(SectionBreakBlock(title=project.name))
        all_blocks.append(HeadingBlock(text=project.name, level=1))
        all_blocks.append(
            ParagraphBlock(runs=[TextRun(text=f"Autor: {project.author}", bold=True)])
        )
        all_blocks.append(
            ParagraphBlock(runs=[TextRun(text=f"Versão: {project.version}")])
        )
        all_blocks.append(
            ParagraphBlock(runs=[TextRun(text=f"Idioma: {project.language}")])
        )
        if project.logo_path and template.logo_enabled:
            all_blocks.append(
                ImageBlock(
                    image_path=str(base_path / project.logo_path),
                    caption="Logo Corporativa",
                )
            )

        all_blocks.append(PageBreakBlock())

        # 2. Gerar Estrutura de Sumário (TOC) preliminar
        toc_blocks = self._generate_toc(project)
        all_blocks.extend(toc_blocks)
        all_blocks.append(PageBreakBlock())

        # 3. Compilar os documentos lendo a estrutura hierárquica
        for vol in project.volumes:
            all_blocks.append(SectionBreakBlock(title=vol.title))
            all_blocks.append(HeadingBlock(text=vol.title, level=1))
            all_blocks.append(PageBreakBlock())

            for part in vol.parts:
                all_blocks.append(HeadingBlock(text=part.title, level=2))
                all_blocks.append(PageBreakBlock())

                for cap in part.chapters:
                    all_blocks.append(HeadingBlock(text=cap.title, level=3))

                    for doc in cap.documents:
                        doc_absolute_path = base_path / doc.file_path
                        if not doc_absolute_path.exists():
                            # Se o arquivo não existir fisicamente, insere um aviso na AST
                            all_blocks.append(
                                ParagraphBlock(
                                    runs=[
                                        TextRun(
                                            text=f"[ERRO: Arquivo de documento ausente: {doc.file_path}]",
                                            bold=True,
                                        )
                                    ]
                                )
                            )
                            continue

                        doc_blocks = self._parse_file_to_blocks(
                            doc_absolute_path, base_path
                        )
                        all_blocks.extend(doc_blocks)

                    # Quebra de página no fim do capítulo
                    all_blocks.append(PageBreakBlock())

        # 4. Salva a AST gerada em um arquivo temporário JSON
        temp_file = tempfile.NamedTemporaryFile(
            suffix=".json", delete=False, mode="w", encoding="utf-8"
        )
        temp_file_path = Path(temp_file.name)

        json_ast = ASTSerializer.serialize_to_json(all_blocks)
        temp_file.write(json_ast)
        temp_file.close()

        return temp_file_path

    def _read_file_resilient(self, file_path: Path) -> str:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            with open(file_path, "r", encoding="latin-1") as f:
                return f.read()

    def _generate_toc(self, project: Project) -> List[Block]:
        """Gera a árvore estruturada do sumário como blocos AST."""
        blocks: List[Block] = []
        blocks.append(HeadingBlock(text="Sumário", level=2))
        blocks.append(
            ParagraphBlock(
                runs=[
                    TextRun(text="Estrutura organizacional do documento:", italic=True)
                ]
            )
        )

        for vol in project.volumes:
            blocks.append(ParagraphBlock(runs=[TextRun(text=vol.title, bold=True)]))
            for part in vol.parts:
                blocks.append(
                    ParagraphBlock(
                        runs=[
                            TextRun(text="    "),
                            TextRun(text=part.title, bold=True, italic=True),
                        ]
                    )
                )
                for cap in part.chapters:
                    blocks.append(
                        ParagraphBlock(
                            runs=[TextRun(text="        "), TextRun(text=cap.title)]
                        )
                    )
        return blocks

    def _parse_file_to_blocks(self, file_path: Path, base_path: Path) -> List[Block]:
        ext = file_path.suffix.lower()
        if ext in [".md", ".markdown"]:
            return self._parse_markdown(file_path, base_path)
        elif ext in [".html", ".htm"]:
            return self._parse_html(file_path, base_path)
        elif ext == ".txt":
            return self._parse_txt(file_path)
        elif ext == ".docx":
            return self._parse_docx(file_path)
        elif ext == ".odt":
            return self._parse_odt(file_path, base_path)

        return [ParagraphBlock(runs=[TextRun(text=f"[Formato não suportado: {ext}]")])]

    def _parse_markdown(self, file_path: Path, base_path: Path) -> List[Block]:
        try:
            md_text = self._read_file_resilient(file_path)
            # Converte Markdown para HTML (com suporte a tabelas e atributos)
            html_text = markdown.markdown(
                md_text, extensions=["tables", "fenced_code", "attr_list"]
            )
            return self._parse_html_string(html_text, base_path)
        except Exception as e:
            return [
                ParagraphBlock(
                    runs=[
                        TextRun(
                            text=f"[Falha ao ler Markdown {file_path.name}: {e}]",
                            bold=True,
                        )
                    ]
                )
            ]

    def _parse_html(self, file_path: Path, base_path: Path) -> List[Block]:
        try:
            html_content = self._read_file_resilient(file_path)
            return self._parse_html_string(html_content, base_path)
        except Exception as e:
            return [
                ParagraphBlock(
                    runs=[
                        TextRun(
                            text=f"[Falha ao ler HTML {file_path.name}: {e}]", bold=True
                        )
                    ]
                )
            ]

    def _parse_html_string(self, html_content: str, base_path: Path) -> List[Block]:
        blocks: List[Block] = []
        soup = BeautifulSoup(html_content, "html.parser")

        for elem in soup.children:
            if elem.name is None:
                continue

            if elem.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
                level = (
                    int(elem.name[1]) + 2
                )  # Desloca os headings do documento para não conflitar com Títulos de Capítulos
                blocks.append(HeadingBlock(text=elem.get_text(), level=level))

            elif elem.name == "p":
                runs = self._parse_html_spans(elem)
                blocks.append(ParagraphBlock(runs=runs))

            elif elem.name in ["ul", "ol"]:
                is_ordered = elem.name == "ol"
                list_block = ListBlock(ordered=is_ordered)
                for li in elem.find_all("li", recursive=False):
                    list_block.items.append(ListItem(runs=self._parse_html_spans(li)))
                blocks.append(list_block)

            elif elem.name == "table":
                table_block = TableBlock()
                # Headers
                thead = elem.find("thead")
                if thead:
                    th_elems = thead.find_all("th")
                    table_block.headers = [
                        TableCell(runs=self._parse_html_spans(th)) for th in th_elems
                    ]
                else:
                    first_tr = elem.find("tr")
                    if first_tr:
                        th_elems = first_tr.find_all(["th", "td"])
                        table_block.headers = [
                            TableCell(runs=self._parse_html_spans(th))
                            for th in th_elems
                        ]

                # Rows
                tbody = elem.find("tbody") or elem
                rows = tbody.find_all("tr")
                # Se não havia thead, pula o primeiro tr que usamos como header
                start_row = 1 if (not thead and len(rows) > 0) else 0
                for tr in rows[start_row:]:
                    row_cells = []
                    for td in tr.find_all(["td", "th"]):
                        row_cells.append(TableCell(runs=self._parse_html_spans(td)))
                    table_block.rows.append(row_cells)
                blocks.append(table_block)

            elif elem.name == "img":
                src = elem.get("src", "")
                alt = elem.get("alt", "Imagem")
                # Se for local, garante o caminho correto
                img_path = src
                if not src.startswith(("http://", "https://", "data:")):
                    img_path = str(base_path / src)
                blocks.append(ImageBlock(image_path=img_path, caption=alt))

            elif elem.name in ["div", "section", "article"]:
                # Chamada recursiva para divs aninhadas
                blocks.extend(self._parse_html_string(str(elem), base_path))

        return blocks

    def _parse_html_spans(self, element) -> List[TextRun]:
        runs: List[TextRun] = []
        for child in element.children:
            if child.name is None:  # Texto simples
                if child.strip():
                    runs.append(TextRun(text=child))
            elif child.name == "strong" or child.name == "b":
                runs.append(TextRun(text=child.get_text(), bold=True))
            elif child.name == "em" or child.name == "i":
                runs.append(TextRun(text=child.get_text(), italic=True))
            elif child.name == "u":
                runs.append(TextRun(text=child.get_text(), underline=True))
            elif child.name == "a":
                href = child.get("href", "")
                runs.append(TextRun(text=child.get_text(), link_url=href))
            else:
                # Trata qualquer outra tag interna extraindo seu texto puro
                text = child.get_text()
                if text.strip():
                    runs.append(TextRun(text=text))
        return runs

    def _parse_txt(self, file_path: Path) -> List[Block]:
        blocks: List[Block] = []
        try:
            content = self._read_file_resilient(file_path)
            lines = content.splitlines()
            for line in lines:
                stripped = line.strip()
                if stripped:
                    blocks.append(ParagraphBlock(runs=[TextRun(text=stripped)]))
        except Exception as e:
            blocks.append(
                ParagraphBlock(
                    runs=[TextRun(text=f"[Erro ao ler arquivo texto: {e}]", bold=True)]
                )
            )
        return blocks

    def _parse_docx(self, file_path: Path) -> List[Block]:
        blocks: List[Block] = []
        try:
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue

                # Mapeia headings do DOCX para headings de nível correspondente
                if paragraph.style.name.startswith("Heading"):
                    try:
                        level = int(paragraph.style.name.replace("Heading ", "")) + 2
                    except ValueError:
                        level = 3
                    blocks.append(HeadingBlock(text=text, level=level))
                else:
                    # Converte runs formatadas
                    runs = []
                    for run in paragraph.runs:
                        runs.append(
                            TextRun(
                                text=run.text,
                                bold=bool(run.bold),
                                italic=bool(run.italic),
                                underline=bool(run.underline),
                            )
                        )
                    blocks.append(ParagraphBlock(runs=runs))

            # Converte tabelas do DOCX
            for table in doc.tables:
                table_block = TableBlock()
                if len(table.rows) == 0:
                    continue
                # Primeiro tr como header
                header_row = table.rows[0]
                table_block.headers = [
                    TableCell(runs=[TextRun(text=cell.text)])
                    for cell in header_row.cells
                ]

                for row in table.rows[1:]:
                    row_cells = []
                    for cell in row.cells:
                        row_cells.append(TableCell(runs=[TextRun(text=cell.text)]))
                    table_block.rows.append(row_cells)
                blocks.append(table_block)

        except Exception as e:
            blocks.append(
                ParagraphBlock(
                    runs=[
                        TextRun(
                            text=f"[Erro ao ler DOCX {file_path.name}: {e}]", bold=True
                        )
                    ]
                )
            )
        return blocks

    def _parse_odt(self, file_path: Path, base_path: Path) -> List[Block]:
        """
        Lê arquivos ODT chamando o LibreOffice headless para exportar temporariamente para HTML,
        depois processa o HTML resultante.
        """
        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                cmd = [
                    "libreoffice",
                    "--headless",
                    "--convert-to",
                    "html",
                    "--outdir",
                    temp_dir,
                    str(file_path),
                ]
                subprocess.run(
                    cmd,
                    stdout=subprocess.DEVNULL,
                    stderr=subprocess.DEVNULL,
                    check=True,
                )

                html_temp_file = Path(temp_dir) / f"{file_path.stem}.html"
                if html_temp_file.exists():
                    return self._parse_html(html_temp_file, base_path)
        except Exception as e:
            return [
                ParagraphBlock(
                    runs=[
                        TextRun(
                            text=f"[Erro ao converter e ler ODT {file_path.name}: {e}]",
                            bold=True,
                        )
                    ]
                )
            ]

        return [
            ParagraphBlock(
                runs=[TextRun(text="[Erro ao processar ODT: conversão falhou.]")]
            )
        ]
